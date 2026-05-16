"""DOCX loader using `python-docx`."""
from __future__ import annotations

import os
from typing import Any

from utils.hashing import sha256_file
from utils.logger import get_logger
from utils.text_cleaning import clean_text

from .base import IngestedDocument, IngestedSection

logger = get_logger(__name__)


def _heading_level(style_name: str | None) -> int:
    """Return heading level for a paragraph style; 0 if not a heading."""
    if not style_name:
        return 0
    s = style_name.lower()
    if s.startswith("heading "):
        try:
            return int(s.replace("heading ", "").strip())
        except ValueError:
            return 1
    if s == "title":
        return 1
    return 0


class DocxLoader:
    name = "docx"

    def supports(self, location: str, mimetype: str | None = None) -> bool:
        return location.lower().endswith(".docx") or mimetype == (
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

    def load(self, location: str, **_: Any) -> IngestedDocument:
        from docx import Document  # type: ignore

        doc = Document(location)
        sections: list[IngestedSection] = []
        current_heading = ""
        current_level = 0
        buffer: list[str] = []

        def flush() -> None:
            if buffer:
                sections.append(
                    IngestedSection(
                        text=clean_text("\n".join(buffer)),
                        heading=current_heading,
                        level=current_level,
                    )
                )
                buffer.clear()

        for para in doc.paragraphs:
            text = (para.text or "").strip()
            if not text:
                continue
            level = _heading_level(para.style.name if para.style else None)
            if level:
                flush()
                current_heading = text
                current_level = level
            else:
                buffer.append(text)
        flush()

        title = (
            doc.core_properties.title
            or os.path.basename(location)
        )

        return IngestedDocument(
            source_type="docx",
            title=title,
            location=os.path.abspath(location),
            checksum=sha256_file(location),
            sections=sections,
            metadata={
                "author": doc.core_properties.author or "",
                "subject": doc.core_properties.subject or "",
            },
        )
